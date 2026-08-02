"""
==========================================================
DOCX Extractor
==========================================================

Responsibilities
----------------
• Read Word document
• Extract text
• Extract tables
• Extract document metadata

No AI logic.
"""

from pathlib import Path

from docx import Document


def extract_docx(
    file_path: str
) -> dict:
    """
    Extract DOCX document.

    Parameters
    ----------
    file_path : str

    Returns
    -------
    dict
    """

    path = Path(file_path)

    if not path.exists():

        raise FileNotFoundError(

            f"DOCX file not found: {path}"

        )

    try:

        document = Document(

            path

        )

    except Exception as exc:

        raise RuntimeError(

            f"Unable to open DOCX: {exc}"

        ) from exc

    # -----------------------------------------------------
    # Extract Paragraphs
    # -----------------------------------------------------

    paragraphs = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            paragraphs.append(

                text

            )

    # -----------------------------------------------------
    # Extract Tables
    # -----------------------------------------------------

    tables = []

    for table in document.tables:

        rows = []

        for row in table.rows:

            rows.append(

                [

                    cell.text.strip()

                    for cell in row.cells

                ]

            )

        tables.append(

            rows

        )

    # -----------------------------------------------------
    # Document Properties
    # -----------------------------------------------------

    properties = document.core_properties

    # -----------------------------------------------------
    # Return
    # -----------------------------------------------------

    return {

        "content": "\n".join(

            paragraphs

        ),

        "tables": tables,

        "metadata": {

            "file_name": path.name,

            "document_type": "DOCX",

            "title": properties.title,

            "author": properties.author,

            "subject": properties.subject,

            "category": properties.category,

            "created": (

                str(properties.created)

                if properties.created

                else None

            ),

            "modified": (

                str(properties.modified)

                if properties.modified

                else None

            ),

            "paragraphs": len(

                paragraphs

            ),

            "tables": len(

                tables

            )

        }

    }